# -*- coding: utf-8 -*-
"""从八字命理目录下的 .docx 文件中提取全文，输出为 Markdown 文件。

用法：
  python extract_docx.py          # 提取所有 .docx 到 knowledge-base/04-古籍原文/
  python extract_docx.py --list   # 仅列出文件信息，不提取
  python extract_docx.py --file "穷通宝鉴-AI知识库.docx"  # 仅提取指定文件
"""
import os, sys, re, hashlib
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SOURCE_DIR = os.path.normpath(os.path.join(ROOT, '..', '八字命理'))
OUT_DIR = os.path.join(ROOT, 'knowledge-base', '03-古籍原文')

DOCX_MAP = {
    '穷通宝鉴-AI知识库.docx': '穷通宝鉴-AI知识库_全文.md',
    '《三命通会》完整知识梳理.docx': '三命通会_完整知识梳理.md',
    '《渊海子平》完整知识梳理.docx': '渊海子平_完整知识梳理.md',
    '《滴天髓》完整知识梳理.docx': '滴天髓_完整知识梳理.md',
    '生辰八字与五行对照表.docx': '生辰八字与五行对照表_全文.md',
    '算命流程.docx': '算命流程_全文.md',
}

def extract_paragraphs(doc):
    """提取段落文本，保留标题层级"""
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        style = para.style.name if para.style else ''
        # 根据样式判断标题级别
        if 'Heading 1' in style or 'heading 1' in style or '标题 1' in style:
            lines.append(f'\n# {text}\n')
        elif 'Heading 2' in style or 'heading 2' in style or '标题 2' in style:
            lines.append(f'\n## {text}\n')
        elif 'Heading 3' in style or 'heading 3' in style or '标题 3' in style:
            lines.append(f'\n### {text}\n')
        elif 'Heading' in style or '标题' in style:
            lines.append(f'\n#### {text}\n')
        elif 'List' in style:
            lines.append(f'- {text}')
        else:
            # 检测常见标题模式
            if re.match(r'^第[一二三四五六七八九十百千]+[章节篇]', text):
                lines.append(f'\n## {text}\n')
            elif re.match(r'^[一二三四五六七八九十]、', text):
                lines.append(f'\n### {text}\n')
            elif re.match(r'^（[一二三四五六七八九十]）', text):
                lines.append(f'\n#### {text}\n')
            else:
                lines.append(text)

    return '\n'.join(lines)

def extract_tables(doc):
    """提取表格为 Markdown 表格"""
    result = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue

        result.append(f'\n### 表格 {i+1}\n')
        # 表头 + 分隔线
        result.append('| ' + ' | '.join(rows[0]) + ' |')
        result.append('| ' + ' | '.join(['---'] * len(rows[0])) + ' |')
        for row in rows[1:]:
            result.append('| ' + ' | '.join(row) + ' |')
        result.append('')

    return '\n'.join(result)

def extract_docx(docx_path, out_path):
    """提取单个 docx 文件"""
    print(f'[提取] {os.path.basename(docx_path)} ...')
    doc = Document(docx_path)

    content = []
    content.append(f'# {os.path.splitext(os.path.basename(out_path))[0]}\n')
    content.append(f'> 来源：{os.path.basename(docx_path)}')
    content.append(f'> 段落数：{len(doc.paragraphs)}，表格数：{len(doc.tables)}')
    content.append(f'> 自动提取时间：{__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")}\n')

    # 提取正文
    content.append('\n---\n')
    content.append('## 正文\n')
    content.append(extract_paragraphs(doc))

    # 提取表格
    content.append('\n---\n')
    content.append('## 表格\n')
    tables_md = extract_tables(doc)
    if tables_md.strip():
        content.append(tables_md)
    else:
        content.append('（无表格）\n')

    full = '\n'.join(content)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(full)

    size = len(full)
    print(f'  → 输出：{out_path} ({size:,} chars)')
    return size

def main():
    if '--list' in sys.argv:
        print(f'源目录：{SOURCE_DIR}\n')
        for fname in sorted(os.listdir(SOURCE_DIR)):
            if fname.endswith('.docx'):
                fpath = os.path.join(SOURCE_DIR, fname)
                size = os.path.getsize(fpath)
                print(f'  {fname}  ({size:,} bytes)')
        return

    target = None
    if '--file' in sys.argv:
        idx = sys.argv.index('--file')
        target = sys.argv[idx + 1]

    os.makedirs(OUT_DIR, exist_ok=True)
    total_chars = 0

    for docx_name, out_name in DOCX_MAP.items():
        if target and docx_name != target:
            continue
        docx_path = os.path.join(SOURCE_DIR, docx_name)
        if not os.path.exists(docx_path):
            print(f'[跳过] 文件不存在：{docx_path}')
            continue
        out_path = os.path.join(OUT_DIR, out_name)
        chars = extract_docx(docx_path, out_path)
        total_chars += chars

    print(f'\n✅ 完成！总计 {total_chars:,} 字符')

if __name__ == '__main__':
    main()
